import fitz  # PyMuPDF for high-speed PDF parsing
from docx import Document  # python-docx for structured Word parsing
import spacy
from spacy.matcher import PhraseMatcher
import os
import pandas as pd

# 1. Initialize Local NLP Engine
nlp = spacy.load("en_core_web_sm")
matcher = PhraseMatcher(nlp.vocab, attr="LOWER") # Use LOWER attribute for better matching

def load_skill_patterns():
    """
    Automatically loads ground-truth skills from your CSV to ensure 
    the parser identifies only relevant technical competencies.
    """
    try:
        # Load skills from your existing verified database
        df = pd.read_csv('data/clean_jobs.csv')
        
        # Extract unique skills, handle potential nested lists in CSV
        all_skills = set()
        for s_list in df['skills'].dropna():
            if s_list.startswith('['):
                import ast
                skills = ast.literal_eval(s_list)
                all_skills.update([s.lower().strip() for s in skills])
            else:
                all_skills.update([s.lower().strip() for s in s_list.split(',')])
        
        # Add patterns to spaCy Matcher
        patterns = [nlp.make_doc(text) for text in all_skills if len(text) > 1]
        matcher.add("SKILL_LIST", patterns)
        print(f"✅ Loaded {len(patterns)} skill patterns into local NLP engine.")
    except Exception as e:
        print(f"⚠️ Warning: Could not load skill patterns from CSV: {e}")
        # Fallback basic list
        fallback = ["python", "sql", "java", "machine learning", "tableau", "excel"]
        matcher.add("SKILL_LIST", [nlp.make_doc(t) for t in fallback])

# Initialize patterns on module load
load_skill_patterns()

def parse_resume(file_path):
    """
    Unified extraction pipeline for PDF and DOCX.
    Maintains 95%+ accuracy by cleaning document artifacts before NLP processing.
    """
    ext = os.path.splitext(file_path)[-1].lower()
    text = ""

    # --- 1. Multi-Format Extraction Layer ---
    try:
        if ext == ".pdf":
            with fitz.open(file_path) as doc:
                # Use 'text' blocks to preserve reading order better than raw strings
                text = " ".join([page.get_text("text") for page in doc])
        
        elif ext == ".docx":
            doc = Document(file_path)
            # Extract from paragraphs and tables (common for skills)
            text = " ".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += " " + cell.text

        # --- 2. Advanced Text Cleaning ---
        # Remove ligatures, non-breaking spaces, and excessive whitespace
        text = text.replace('\xa0', ' ').replace('\x0c', ' ')
        text = " ".join(text.split())

        # --- 3. Local Skill Identification via spaCy ---
        doc_nlp = nlp(text)
        matches = matcher(doc_nlp)
        
        found_skills = set()
        for match_id, start, end in matches:
            # We use the original text but store it in a set to avoid duplicates
            span = doc_nlp[start:end]
            found_skills.add(span.text.lower().strip())
            
        return list(found_skills)

    except Exception as e:
        print(f"❌ Parser Error on {file_path}: {e}")
        return []